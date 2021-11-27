from flask import Flask, render_template, Response, request, redirect
from pymysql import connect, cursors, Error
from datetime import datetime  
from docx import Document
from docx.shared import Inches
db_config = {
    "host": "127.0.0.1",
    "user": "root",
    "password": "",
    "database": "blogs",
    "cursorclass": cursors.DictCursor
}
conn = connect(**db_config)
cur=conn.cursor()

app = Flask(__name__)

@app.route('/', methods=["GET"])
def render_form():
    cur = conn.cursor()
    sql="SELECT * FROM posts"
    cur.execute(sql)
    return render_template("index.html",rows=cur)


@app.route('/about')
def about():
    cur = conn.cursor()
    sql="SELECT * from posts"
    cur.execute(sql)
    return render_template("about.html")



@app.route('/index', methods=["POST"])
def new_post():
    cur = conn.cursor()
    title=request.form["post_title"]
    detail=request.form["post_detail"]
    sql="INSERT INTO posts(post_title,post_detail,post_time) VALUES (%s,%s,%s)"
    value=(title,detail,datetime.now())
    try:
        cur.execute(sql,value)
        conn.commit()
    except:
        conn.rollback()
    return redirect("/index", code=302)

@app.route('/edit/<id>', methods=["GET"])
def render_form_edit(id):
    cur = conn.cursor()
    sql="SELECT * from posts where id="+str(id)
    cur.execute(sql)
    for row in cur:
         r0=row[0]
         r1=row[1]
         r2=row[2]

    return render_template("edit.html",r0=r0,r1=r1,r2=r2)



@app.route('/edit/<id>', methods=["POST"])
def edit_post(id):
    cur = conn.cursor()
    title=request.form["post_title"]
    detail=request.form["post_detail"]
    sql=f" UPDATE posts SET post_title='{title}',post_detail='{detail}' where id = {id}"
    try:
        cur.execute(sql)
        conn.commit()
    except:
        conn.rollback()

    return redirect("/edit/"+id, code=302)


@app.route('/letter', methods=["GET"])
def xport_file():

    return render_template("letter.html")

@app.route('/letter', methods=["POST"])
def get_text():
    name=request.form["name"]
    reason=request.form["reason"]
    document = Document()
    document.add_heading(name, 0)
    document.add_paragraph(reason)
    document.save('static/letter.docx')
    return redirect("/letter", code=302)


    

if __name__ == "__main__":
    app.run(debug=True)
